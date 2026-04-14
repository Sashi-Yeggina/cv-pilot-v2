"""
bullet_library.py
─────────────────
Bullet-point level CV library engine.

Core idea:
  Every bullet point from every CV the user has ever uploaded or generated
  is stored and tagged with skills / keywords.  When a new JD arrives we
  search this library first.  Only if we can't cover the JD requirements
  do we call the Claude API — and even then, only for the *gaps*.

Flow
────
1. index_cv()        — called whenever a CV is uploaded / generated.
                       Extracts every bullet, tags it with skills, saves
                       to cv_bullet_library table.

2. match_jd()        — called at the start of every generation request.
                       Extracts JD skills (no API), searches the library,
                       returns the best matching bullets per skill plus
                       a coverage score.

3. assemble_cv()     — builds a complete CV dict from library bullets.
                       Called when coverage ≥ threshold (default 70 %).

4. gap_skills()      — returns the skills NOT covered by the library.
                       Passed to Claude so it only generates the missing
                       sections → fewer tokens, lower cost.
"""

import re
from typing import Any

# ─────────────────────────────────────────────────────────────────────────────
# SKILL TAXONOMY
# A curated keyword list covering cloud / DevOps / infrastructure domains.
# Add more rows as needed — no API calls required.
# ─────────────────────────────────────────────────────────────────────────────
SKILL_TAXONOMY: dict[str, list[str]] = {
    # AWS core
    "ec2":           ["ec2", "elastic compute", "virtual machine", "vm"],
    "s3":            ["s3", "simple storage", "object storage"],
    "rds":           ["rds", "relational database", "aurora", "postgresql", "mysql"],
    "vpc":           ["vpc", "virtual private cloud", "subnet", "route table", "nat gateway"],
    "iam":           ["iam", "identity access", "role policy", "rbac", "least privilege"],
    "lambda":        ["lambda", "serverless", "function as a service", "faas"],
    "cloudwatch":    ["cloudwatch", "cloud watch", "metrics", "alarms", "logs"],
    "cloudfront":    ["cloudfront", "cdn", "content delivery", "edge"],
    "eks":           ["eks", "elastic kubernetes", "kubernetes", "k8s"],
    "ecs":           ["ecs", "elastic container", "fargate"],
    "elb":           ["elb", "elastic load balancer", "alb", "nlb", "load balanc"],
    "auto_scaling":  ["auto scaling", "autoscaling", "asg", "horizontal scaling"],
    "cloudformation":["cloudformation", "cloud formation", "cfn"],
    "route53":       ["route53", "route 53", "dns", "hosted zone"],
    "kms":           ["kms", "key management", "encryption", "encrypt"],
    "guardduty":     ["guardduty", "guard duty", "threat detection"],
    "control_tower": ["control tower", "controltower", "landing zone"],
    "organizations": ["aws organizations", "multi-account", "multi account", "scp"],
    "outposts":      ["outposts", "aws outpost", "hybrid cloud", "on-premises"],
    # IaC
    "terraform":     ["terraform", "hcl", "hashicorp"],
    "ansible":       ["ansible", "playbook", "configuration management"],
    "bicep":         ["bicep", "arm template", "azure resource manager"],
    # Containers
    "docker":        ["docker", "container", "dockerfile", "image"],
    "helm":          ["helm", "helm chart", "chart"],
    "kubernetes":    ["kubernetes", "k8s", "kubectl", "cluster"],
    # CI/CD
    "github_actions":["github actions", "gha", "workflow yml"],
    "jenkins":       ["jenkins", "jenkinsfile", "pipeline"],
    "gitlab_ci":     ["gitlab ci", "gitlab-ci", ".gitlab-ci"],
    # Observability
    "prometheus":    ["prometheus", "prom", "alertmanager"],
    "grafana":       ["grafana", "dashboard", "visualiz"],
    "elk":           ["elk", "elasticsearch", "logstash", "kibana", "opensearch"],
    "splunk":        ["splunk"],
    # Scripting
    "python":        ["python", ".py", "boto3", "flask", "django", "fastapi"],
    "bash":          ["bash", "shell script", "sh script", "bash script"],
    "powershell":    ["powershell", "pwsh", "ps1"],
    # Architecture
    "well_architected": ["well-architected", "well architected", "waf", "five pillars"],
    "landing_zone":  ["landing zone", "landingzone", "account vending"],
    "ha_dr":         ["high availab", "disaster recover", "rpo", "rto", "fault toleran"],
    "cost_optim":    ["cost optim", "right-siz", "rightsiz", "reserved instance", "savings plan"],
    "security":      ["security", "complian", "audit", "sox", "pci", "hipaa", "gdpr"],
    "networking":    ["networking", "network architect", "bgp", "transit gateway", "vpn"],
    # Clouds
    "azure":         ["azure", "microsoft cloud", "aks", "arm"],
    "gcp":           ["gcp", "google cloud", "gke"],
    # Leadership
    "mentoring":     ["mentor", "coach", "upskill", "knowledge transfer"],
    "governance":    ["governance", "guardrail", "policy", "standard", "framework"],
    "migration":     ["migrat", "rehost", "replatform", "refactor", "moderniz"],
    "5g":            ["5g", "telecom", "ran", "edge computing"],
    "devops":        ["devops", "dev ops", "site reliab", "sre"],
}

def extract_skills_from_text(text: str) -> list[str]:
    """
    Return a deduplicated list of canonical skill keys found in `text`.
    Pure keyword match — zero API calls.
    """
    lower = text.lower()
    found: set[str] = set()
    for skill_key, patterns in SKILL_TAXONOMY.items():
        if any(pat in lower for pat in patterns):
            found.add(skill_key)
    return sorted(found)

def extract_keywords_from_text(text: str) -> list[str]:
    """
    Return significant words (len > 3, not stopwords) from text.
    Used as a secondary fuzzy-match signal.
    """
    stopwords = {
        "with", "that", "have", "from", "this", "were", "they",
        "been", "their", "will", "when", "than", "more", "also",
        "into", "each", "over", "such", "while", "across", "using",
        "within", "ensure", "including", "enabling", "ensuring",
    }
    tokens = re.findall(r"[a-z][a-z0-9\-]{2,}", text.lower())
    return [t for t in tokens if t not in stopwords]

# ─────────────────────────────────────────────────────────────────────────────
# DOCX → BULLET EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

def extract_bullets_from_docx(docx_bytes: bytes, cv_meta: dict) -> list[dict]:
    """
    Parse a DOCX file and return a list of bullet dicts ready to insert
    into cv_bullet_library.

    cv_meta keys expected: cv_id, user_id, role_title, seniority
    """
    from docx import Document
    from io import BytesIO

    doc = Document(BytesIO(docx_bytes))
    bullets: list[dict] = []

    current_section = "summary"          # track which section we're in
    current_company = ""

    section_markers = {
        "summary":        ["summary", "profile", "objective"],
        "experience":     ["experience", "employment", "career", "history"],
        "skills":         ["skill", "technical", "competenc", "technolog"],
        "education":      ["education", "degree", "universit", "qualification"],
        "certifications": ["certif", "accredit", "award"],
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 15:        # skip blank / header-only lines
            continue

        # Detect section changes from headings
        style_name = (para.style.name or "").lower()
        is_heading = "heading" in style_name or style_name in ("title", "subtitle")

        if is_heading:
            lower_text = text.lower()
            for sec, markers in section_markers.items():
                if any(m in lower_text for m in markers):
                    current_section = sec
                    break
            # Check if this looks like a company/role line (bold, not a heading style)
            continue

        # Detect company/role lines (bold paragraphs that aren't headings)
        if para.runs and all(r.bold for r in para.runs if r.text.strip()):
            if len(text) < 120 and "|" in text:
                current_company = text.split("|")[0].strip()
            continue

        # Everything else is a content bullet or sentence
        # Skip very long paragraphs (> 500 chars) — probably pasted blobs
        if len(text) > 500:
            continue

        skills   = extract_skills_from_text(text)
        keywords = extract_keywords_from_text(text)

        bullets.append({
            "user_id":         cv_meta["user_id"],
            "cv_id":           cv_meta["cv_id"],
            "text":            text,
            "section":         current_section,
            "skills":          skills,
            "keywords":        keywords[:40],          # cap array size
            "role_context":    cv_meta.get("role_title", ""),
            "company_context": current_company,
            "seniority":       cv_meta.get("seniority", ""),
            "quality_score":   0.6 if current_section == "experience" else 0.4,
        })

    return bullets

# ─────────────────────────────────────────────────────────────────────────────
# JD SKILL EXTRACTION  (no API — keyword match only)
# ─────────────────────────────────────────────────────────────────────────────

def extract_jd_skills_locally(jd_text: str) -> list[str]:
    """
    Extract required skills from a JD using keyword matching.
    Returns canonical skill keys.  No API call.
    """
    return extract_skills_from_text(jd_text)

# ─────────────────────────────────────────────────────────────────────────────
# LIBRARY MATCHING
# ─────────────────────────────────────────────────────────────────────────────

def score_bullet_for_skill(bullet: dict, skill_key: str, jd_keywords: list[str]) -> float:
    """
    Score how well a single library bullet covers a required skill.
    Returns 0.0 – 1.0.
    """
    score = 0.0

    # Primary: does the bullet's skills array contain the skill?
    if skill_key in bullet.get("skills", []):
        score += 0.70

    # Secondary: keyword overlap with JD keywords
    bullet_kws = set(bullet.get("keywords", []))
    jd_kws     = set(jd_keywords)
    if bullet_kws and jd_kws:
        overlap = len(bullet_kws & jd_kws) / max(len(jd_kws), 1)
        score += overlap * 0.20

    # Tertiary: quality score (usage, user feedback)
    score += bullet.get("quality_score", 0.5) * 0.10

    return round(min(score, 1.0), 3)


def match_jd_to_library(
    library_bullets: list[dict],
    jd_skills: list[str],
    jd_text: str,
    top_per_skill: int = 3,
) -> dict:
    """
    Given all library bullets for a user and a list of JD required skills,
    find the best matching bullets for each skill.

    Returns:
    {
      "coverage_pct":    0.82,          # fraction of JD skills covered ≥ threshold
      "covered_skills":  ["eks", ...],
      "gap_skills":      ["control_tower", ...],
      "bullets_by_skill": {
          "eks":   [{"id": ..., "text": ..., "score": 0.91}, ...],
          "iam":   [...],
          ...
      },
      "all_selected_bullets": [...],    # flat deduplicated list, best first
    }
    """
    jd_keywords = extract_keywords_from_text(jd_text)

    bullets_by_skill: dict[str, list[dict]] = {}
    coverage_threshold = 0.50      # a bullet must score ≥ 0.50 to "cover" a skill

    for skill in jd_skills:
        scored: list[tuple[float, dict]] = []
        for b in library_bullets:
            s = score_bullet_for_skill(b, skill, jd_keywords)
            if s >= coverage_threshold:
                scored.append((s, b))

        # Sort descending by score
        scored.sort(key=lambda x: x[0], reverse=True)
        if scored:
            bullets_by_skill[skill] = [
                {**b, "_score": sc}
                for sc, b in scored[:top_per_skill]
            ]

    covered_skills = [sk for sk in jd_skills if sk in bullets_by_skill]
    gap_skills     = [sk for sk in jd_skills if sk not in bullets_by_skill]
    coverage_pct   = len(covered_skills) / max(len(jd_skills), 1)

    # Build a flat deduplicated list of the best bullets (one per skill)
    seen_ids:     set[str] = set()
    all_selected: list[dict] = []
    for skill in covered_skills:
        for b in bullets_by_skill[skill]:
            bid = b.get("id") or b.get("text")[:60]     # fallback key
            if bid not in seen_ids:
                seen_ids.add(bid)
                all_selected.append(b)

    return {
        "coverage_pct":       round(coverage_pct, 3),
        "covered_skills":     covered_skills,
        "gap_skills":         gap_skills,
        "bullets_by_skill":   bullets_by_skill,
        "all_selected_bullets": all_selected,
    }

# ─────────────────────────────────────────────────────────────────────────────
# CV ASSEMBLY FROM LIBRARY
# ─────────────────────────────────────────────────────────────────────────────

def assemble_cv_from_library(
    match_result: dict,
    jd_profile: dict,
    base_cv_text: str,
    api_bullets: list[str] | None = None,
) -> dict:
    """
    Build a CV content dict from library bullets + optional API-generated bullets
    for any gap skills.

    Returns a dict that can be fed directly into the DOCX builder:
    {
      "role_title":    "Senior AWS Architect",
      "summary":       "...",    # best summary bullet from library
      "bullets":       ["...", "...", ...],
      "skills_used":   ["eks", "terraform", ...],
      "gap_skills":    ["control_tower"],
      "source":        "library_only" | "library_plus_api"
    }
    """
    selected = match_result["all_selected_bullets"]
    gap      = match_result["gap_skills"]

    # Pick the best summary bullet (section == 'summary') if available
    summary_bullets = [b for b in selected if b.get("section") == "summary"]
    summary = summary_bullets[0]["text"] if summary_bullets else _build_summary_from_jd(jd_profile)

    # Pick experience bullets (section == 'experience'), ordered by score
    exp_bullets = [b for b in selected if b.get("section") == "experience"]
    exp_bullets.sort(key=lambda b: b.get("_score", 0), reverse=True)
    experience_texts = [b["text"] for b in exp_bullets[:10]]

    # Append API-generated bullets for gap skills (if any)
    if api_bullets:
        experience_texts.extend(api_bullets)

    source = "library_only" if not api_bullets else "library_plus_api"

    return {
        "role_title":  jd_profile.get("role_title", "Cloud Professional"),
        "summary":     summary,
        "bullets":     experience_texts,
        "skills_used": match_result["covered_skills"],
        "gap_skills":  gap,
        "source":      source,
    }


def _build_summary_from_jd(jd_profile: dict) -> str:
    """
    Construct a minimal summary string from JD profile metadata
    when no suitable library summary bullet exists.
    No API call — template interpolation only.
    """
    role    = jd_profile.get("role_title", "Cloud Professional")
    years   = jd_profile.get("years_experience", "several years of")
    skills  = ", ".join(jd_profile.get("required_skills", [])[:4])
    cloud   = jd_profile.get("cloud_platform", "cloud")
    return (
        f"Experienced {role} with {years} experience in {cloud} architecture. "
        f"Strong background in {skills}, with a focus on reliability, security, and automation."
    )

# ─────────────────────────────────────────────────────────────────────────────
# DOCX WRITER  — applies assembled content to the user's template
# ─────────────────────────────────────────────────────────────────────────────

def write_assembled_cv_to_docx(
    template_bytes: bytes,
    assembled: dict,
) -> bytes:
    """
    Write the assembled CV content into the template DOCX file.
    Uses python-docx to find and replace key sections.

    Replacement targets (in order):
      1. First paragraph  → role_title
      2. First non-heading paragraph after PROFESSIONAL SUMMARY → summary
      3. All bullet list paragraphs  → assembled bullets (in order)
    """
    from docx import Document
    from io import BytesIO

    doc = Document(BytesIO(template_bytes))
    paras = doc.paragraphs
    n = len(paras)

    # 1. Update title (paragraph 0)
    if n > 0 and assembled.get("role_title"):
        _set_para_text(paras[0], assembled["role_title"])

    # 2. Find PROFESSIONAL SUMMARY section and update the next content paragraph
    summary_written = False
    for i, para in enumerate(paras):
        if "summary" in para.text.lower() and _is_heading(para):
            # Next non-empty paragraph is the summary body
            for j in range(i + 1, min(i + 5, n)):
                if paras[j].text.strip() and not _is_heading(paras[j]):
                    _set_para_text(paras[j], assembled.get("summary", ""))
                    summary_written = True
                    break
            break

    # 3. Replace bullet list paragraphs with assembled bullets
    bullet_pool = list(assembled.get("bullets", []))   # mutable copy
    bullets_written = 0
    for para in paras:
        style = (para.style.name or "").lower()
        is_list_para = (
            "list" in style
            or "bullet" in style
            or para._element.pPr is not None and para._element.pPr.numPr is not None
        )
        if is_list_para and para.text.strip():
            if bullet_pool:
                _set_para_text(para, bullet_pool.pop(0))
                bullets_written += 1
            else:
                # No more replacement bullets — clear the placeholder
                _set_para_text(para, "")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _is_heading(para: Any) -> bool:
    style = (para.style.name or "").lower()
    return "heading" in style or style in ("title", "subtitle")


def _set_para_text(para: Any, text: str) -> None:
    """Replace all runs in a paragraph with a single run containing `text`."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

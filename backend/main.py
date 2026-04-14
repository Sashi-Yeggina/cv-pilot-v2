"""Main FastAPI application for CV Pilot API"""
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, WebSocket, WebSocketDisconnect, status, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from contextlib import asynccontextmanager
import os
import io
from datetime import datetime
from uuid import uuid4

# Local imports
from config import settings
from models import *
from auth_service import register_user, login_user, verify_token, logout_user
from database import *
from database import (
    get_all_users_usage_stats, get_user_usage_stats,
    check_jd_duplicate,
    create_bulk_job, update_bulk_job_progress, get_user_bulk_jobs,
    create_submission, get_user_submissions, update_submission,
    get_submission, get_submission_pipeline,
)
from cv_engine import (
    parse_jd, score_cvs, enhance_cv, extract_cv_info,
    fill_gap_bullets, decide_generation_strategy,
    LIBRARY_ONLY_THRESHOLD, LIBRARY_PATCH_THRESHOLD,
    AVAILABLE_MODELS, DEFAULT_MODEL,
)
from bullet_library import (
    extract_bullets_from_docx,
    extract_jd_skills_locally,
    match_jd_to_library,
    assemble_cv_from_library,
    write_assembled_cv_to_docx,
)
from storage_service import upload_cv, download_cv, delete_cv, get_file_url
from docx import Document
import json

# WebSocket Connection Manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: dict = {}

    async def connect(self, websocket: WebSocket, user_id: str):
        await websocket.accept()
        self.active_connections[user_id] = websocket

    def disconnect(self, user_id: str):
        if user_id in self.active_connections:
            del self.active_connections[user_id]

    async def broadcast(self, event: str, data: dict):
        """Broadcast event to all connected admin clients"""
        # In production, check if user is admin
        for user_id, connection in self.active_connections.items():
            try:
                await connection.send_json({
                    "event": event,
                    "data": data,
                    "timestamp": datetime.utcnow().isoformat()
                })
            except Exception:
                pass

manager = ConnectionManager()

# ════════════════════════════════════════════════════════════════
# LIFESPAN
# ════════════════════════════════════════════════════════════════

@asynccontextmanager
async def lifespan(app: FastAPI):
    print("CV Pilot API starting...")
    yield
    print("CV Pilot API shutting down...")

# ════════════════════════════════════════════════════════════════
# CREATE FASTAPI APP
# ════════════════════════════════════════════════════════════════

app = FastAPI(
    title="CV Pilot API",
    description="Multi-user CV generation platform",
    version="1.0.0",
    lifespan=lifespan
)

# ════════════════════════════════════════════════════════════════
# CORS MIDDLEWARE
# ════════════════════════════════════════════════════════════════

app.add_middleware(
    CORSMiddleware,
    allow_origins=[settings.FRONTEND_URL],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ════════════════════════════════════════════════════════════════
# HEALTH CHECK
# ════════════════════════════════════════════════════════════════

@app.get("/health", response_model=HealthResponse)
async def health():
    """Health check endpoint"""
    return HealthResponse(status="ok")

@app.get("/")
async def root():
    """Root endpoint"""
    return {"message": "CV Pilot API v1.0.0", "status": "running"}

# ════════════════════════════════════════════════════════════════
# AUTH ENDPOINTS
# ════════════════════════════════════════════════════════════════

@app.post("/api/auth/register", response_model=AuthResponse)
async def register(request: RegisterRequest):
    """Register new user"""
    try:
        result = await register_user(request.email, request.password, request.full_name)

        if "error" in result:
            raise HTTPException(status_code=400, detail=result["error"])

        # Log activity
        create_activity_log(
            user_id=result["user_id"],
            action_type="register",
            description=f"User registered"
        )

        return AuthResponse(
            access_token=result.get("access_token") or "",
            user_id=result["user_id"],
            email=result["email"],
            role="user"
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/api/auth/login", response_model=AuthResponse)
async def login(request: LoginRequest):
    """User login"""
    try:
        result = await login_user(request.email, request.password)

        if "error" in result:
            raise HTTPException(status_code=401, detail="Invalid credentials")

        # Log activity
        create_activity_log(
            user_id=result["user_id"],
            action_type="login",
            description=f"User logged in"
        )

        return AuthResponse(
            access_token=result["access_token"],
            user_id=result["user_id"],
            email=result["email"],
            role=result.get("role", "user")
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid credentials")

@app.post("/api/auth/logout")
async def logout(token: str = Depends(lambda: None)):
    """User logout"""
    result = logout_user(token)
    return result

# ════════════════════════════════════════════════════════════════
# CV ENDPOINTS
# ════════════════════════════════════════════════════════════════

def get_current_user_id(authorization: str = None) -> str:
    """Extract user ID from token"""
    if not authorization:
        raise HTTPException(status_code=401, detail="Not authenticated")

    try:
        token = authorization.replace("Bearer ", "")
        user = verify_token(token)
        if not user:
            raise HTTPException(status_code=401, detail="Invalid token")
        return user.user.id
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid token")

@app.post("/api/cv/upload")
async def upload_cv_file(
    file: UploadFile = File(...),
    cv_type: str = "base",
    authorization: str = None
):
    """Upload CV file"""
    try:
        user_id = get_current_user_id(authorization)

        # Read file
        contents = await file.read()
        file_size = len(contents)

        # Determine bucket
        bucket = "user_cvs" if cv_type in ["base", "template"] else "generated_cvs"

        # Extract text from DOCX
        doc = Document(io.BytesIO(contents))
        cv_text = "\n".join([para.text for para in doc.paragraphs])

        # Extract metadata (use user's assigned model; fall back to default)
        _user_model = get_user_model(user_id, default=DEFAULT_MODEL)
        cv_info = extract_cv_info(cv_text, model=_user_model)["result"]

        # Upload to storage
        file_path = f"{user_id}/{file.filename}"
        storage_path = upload_cv(user_id, bucket, file_path, contents)

        if not storage_path:
            raise Exception("Failed to upload file")

        # Create database record
        cv_record = create_cv(
            user_id=user_id,
            filename=file.filename,
            file_path=storage_path,
            file_size_bytes=file_size,
            cv_type=cv_type,
            role_title=cv_info.get("role_title"),
            seniority=cv_info.get("seniority"),
            industry=cv_info.get("industry")
        )

        # ── Index bullets into library (only for base CVs, not templates) ──────
        bullets_indexed = 0
        if cv_type == "base" and cv_record:
            try:
                bullets = extract_bullets_from_docx(
                    docx_bytes=contents,
                    cv_meta={
                        "cv_id":       cv_record["id"],
                        "user_id":     user_id,
                        "role_context":cv_info.get("role_title"),
                        "seniority":   cv_info.get("seniority"),
                    }
                )
                if bullets:
                    bullets_indexed = save_bullets_for_cv(bullets)
                    print(f"[upload] indexed {bullets_indexed} bullets for cv {cv_record['id']}")
            except Exception as idx_err:
                # Non-fatal: upload succeeds even if indexing fails
                print(f"[upload] bullet indexing failed (non-fatal): {idx_err}")

        # Log activity
        create_activity_log(
            user_id=user_id,
            action_type="cv_upload",
            description=f"Uploaded CV: {file.filename}",
            metadata={"cv_type": cv_type, "file_size": file_size, "bullets_indexed": bullets_indexed}
        )

        # Broadcast to admin
        await manager.broadcast("cv_uploaded", {
            "user_id": user_id,
            "filename": file.filename,
            "cv_type": cv_type,
            "bullets_indexed": bullets_indexed,
        })

        return {
            "id": cv_record["id"],
            "filename": file.filename,
            "cv_type": cv_type,
            "role_title": cv_info.get("role_title"),
            "seniority": cv_info.get("seniority"),
            "industry": cv_info.get("industry"),
            "bullets_indexed": bullets_indexed,
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ─────────────────────────────────────────────────────────────────────────────
# LIBRARY MATCH PREVIEW — lets the UI show coverage before the user clicks Generate
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/api/cv/match")
async def match_cv_to_jd(
    request: dict,
    authorization: str = None
):
    """
    Preview how well the user's bullet library covers a job description.
    Returns coverage percentage, covered skills, gap skills, and a sample
    of matched bullets — all computed locally, no Claude API call needed.

    Request body:
        { "jd_text": "<full job description text>" }
    """
    try:
        user_id = get_current_user_id(authorization)
        jd_text  = request.get("jd_text", "").strip()

        if not jd_text:
            raise HTTPException(status_code=400, detail="jd_text is required")

        # 1. Parse JD for role/skills using the user's assigned model
        user_model = get_user_model(user_id, default=DEFAULT_MODEL)
        jd_profile = parse_jd(jd_text, model=user_model)

        # 2. Extract required skills with zero-cost keyword matching
        jd_skills = extract_jd_skills_locally(jd_text, jd_profile)

        # 3. Load the user's bullet library
        library_bullets = get_user_library_bullets(user_id)

        # 4. Match — pure Python, no API
        match_result = match_jd_to_library(library_bullets, jd_skills, jd_text)

        # 5. Build a trimmed response (don't send full bullet texts in preview)
        covered_preview = [
            {
                "skill": skill,
                "top_bullet": bullets[0]["text"][:120] + "…" if bullets else "",
                "bullet_count": len(bullets),
            }
            for skill, bullets in match_result.get("skill_bullets", {}).items()
            if bullets
        ]

        return {
            "coverage_pct":      round(match_result["coverage_pct"] * 100, 1),
            "strategy":          decide_generation_strategy(match_result["coverage_pct"]),
            "jd_skills_total":   len(jd_skills),
            "covered_skills":    match_result["covered_skills"],
            "gap_skills":        match_result["gap_skills"],
            "library_size":      len(library_bullets),
            "covered_preview":   covered_preview,
            "estimated_api_calls": (
                1 if decide_generation_strategy(match_result["coverage_pct"]) == "library_only" else
                2 if decide_generation_strategy(match_result["coverage_pct"]) == "library_plus_api" else
                3
            ),
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/api/cv/list", response_model=CVListResponse)
async def list_cvs(cv_type: str = None, authorization: str = None):
    """List user's CVs"""
    try:
        user_id = get_current_user_id(authorization)
        cvs = get_user_cvs(user_id, cv_type)

        return CVListResponse(
            cvs=[CVResponse(
                id=cv["id"],
                filename=cv["filename"],
                file_path=cv.get("file_path"),
                cv_type=cv["cv_type"],
                created_at=cv["created_at"],
                role_title=cv.get("role_title"),
                seniority=cv.get("seniority"),
                industry=cv.get("industry")
            ) for cv in cvs],
            count=len(cvs)
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/api/cv/{cv_id}/download")
async def download_cv_file(cv_id: str, authorization: str = None):
    """Download CV file"""
    try:
        user_id = get_current_user_id(authorization)
        cv = get_cv(cv_id, user_id)

        if not cv:
            raise HTTPException(status_code=404, detail="CV not found")

        file_data = download_cv("user_cvs" if cv["cv_type"] in ["base", "template"] else "generated_cvs", cv["file_path"])

        if not file_data:
            raise HTTPException(status_code=404, detail="File not found")

        # Log activity
        create_activity_log(
            user_id=user_id,
            action_type="cv_download",
            description=f"Downloaded CV: {cv['filename']}"
        )

        return FileResponse(
            io.BytesIO(file_data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=cv["filename"]
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.delete("/api/cv/{cv_id}")
async def delete_cv_file(cv_id: str, authorization: str = None):
    """Delete CV"""
    try:
        user_id = get_current_user_id(authorization)
        cv = get_cv(cv_id, user_id)

        if not cv:
            raise HTTPException(status_code=404, detail="CV not found")

        delete_cv("user_cvs" if cv["cv_type"] in ["base", "template"] else "generated_cvs", cv["file_path"])
        delete_cv_db(cv_id, user_id)

        # Log activity
        create_activity_log(
            user_id=user_id,
            action_type="cv_delete",
            description=f"Deleted CV: {cv['filename']}"
        )

        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ════════════════════════════════════════════════════════════════
# JD ENDPOINTS
# ════════════════════════════════════════════════════════════════

@app.post("/api/jd/create", response_model=JDResponse)
async def create_jd_endpoint(request: JDRequest, authorization: str = None):
    """Create job description (saved to DB; generation must succeed for it to persist)."""
    try:
        user_id = get_current_user_id(authorization)

        jd = create_jd(
            user_id=user_id,
            full_text=request.full_text,
            role_title=request.role_title,
            company_name=request.company_name,
            industry=request.industry,
            vendor_name=request.vendor_name,
            client_name=request.client_name,
            client_email=request.client_email,
            notes=request.notes,
        )

        # Log activity
        create_activity_log(
            user_id=user_id,
            action_type="jd_create",
            description=f"Created JD: {request.role_title or 'Unknown role'}",
            metadata={
                "vendor_name": request.vendor_name,
                "client_name": request.client_name,
                "client_email": request.client_email,
            }
        )

        return JDResponse(
            id=jd["id"],
            role_title=jd.get("role_title"),
            company_name=jd.get("company_name"),
            industry=jd.get("industry"),
            vendor_name=jd.get("vendor_name"),
            client_name=jd.get("client_name"),
            client_email=jd.get("client_email"),
            created_at=jd["created_at"]
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/api/jd/list", response_model=JDListResponse)
async def list_jds(authorization: str = None):
    """List user's JDs"""
    try:
        user_id = get_current_user_id(authorization)
        jds = get_user_jds(user_id)

        return JDListResponse(
            jds=[JDResponse(
                id=jd["id"],
                role_title=jd.get("role_title"),
                company_name=jd.get("company_name"),
                industry=jd.get("industry"),
                vendor_name=jd.get("vendor_name"),
                client_name=jd.get("client_name"),
                client_email=jd.get("client_email"),
                created_at=jd["created_at"]
            ) for jd in jds],
            count=len(jds)
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/jd/check-duplicate", response_model=JDDuplicateWarning)
async def check_jd_duplicate_endpoint(request: dict, authorization: str = None):
    """
    Check whether the same JD text + client_email combination was already
    submitted by a different user.  Returns a warning so the UI can show a
    caution banner before the recruiter proceeds.

    Request body:
        { "jd_text": "<full text>", "client_email": "<optional email>" }
    """
    try:
        get_current_user_id(authorization)   # auth check only
        jd_text      = (request.get("jd_text") or "").strip()
        client_email = (request.get("client_email") or "").strip() or None

        if not jd_text:
            raise HTTPException(status_code=400, detail="jd_text is required")

        result = check_jd_duplicate(jd_text, client_email)
        return JDDuplicateWarning(**result)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ════════════════════════════════════════════════════════════════
# GENERATION ENDPOINTS
# ════════════════════════════════════════════════════════════════

@app.post("/api/cv/generate", response_model=GenerationResponse)
async def generate_cv(
    request: GenerationRequest,
    background_tasks: BackgroundTasks,
    authorization: str = None,
):
    """Generate new CV from 1–3 base CVs + JD (async — returns immediately)."""
    try:
        user_id = get_current_user_id(authorization)

        # Create generation record with status "pending"
        gen = create_generation(
            user_id=user_id,
            jd_id=request.jd_id,
            base_cv_ids=request.base_cv_ids,
            template_cv_id=request.template_cv_id,
            status="pending",
        )
        if not gen:
            raise HTTPException(status_code=500, detail="Failed to create generation record")

        # Fire generation in background — request returns immediately
        background_tasks.add_task(
            process_generation_sync,
            gen["id"], user_id, request, authorization,
        )

        await manager.broadcast("generation_started", {
            "user_id": user_id,
            "generation_id": gen["id"],
            "jd_id": request.jd_id,
        })

        return GenerationResponse(
            id=gen["id"],
            status="pending",
            jd_id=gen["jd_id"],
            base_cv_ids=gen["base_cv_ids"],
            created_at=gen["created_at"],
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

def process_generation_sync(generation_id: str, user_id: str, request: GenerationRequest, authorization: str):
    """
    Library-First CV Generation
    ────────────────────────────
    Decision tree (saves API cost):
      coverage ≥ 0.80  →  library_only      0 extra API calls
      coverage ≥ 0.55  →  library_plus_api  1 small call (gap bullets only)
      coverage <  0.55 →  full_api          score_cvs + enhance_cv (2 calls)
    """
    try:
        start_time = datetime.utcnow()

        # ── Step 1: Fetch the job description ────────────────────────────────
        jd = get_jd(request.jd_id, user_id)
        if not jd:
            update_generation(generation_id, "failed", error_message="JD not found")
            return

        # ── Step 2: Fetch the model assigned to this user by admin ──────────
        user_model = get_user_model(user_id, default=DEFAULT_MODEL)
        print(f"[gen:{generation_id}] using model={user_model} for user={user_id}")

        # Accumulators for token tracking across all API calls
        total_input_tokens  = 0
        total_output_tokens = 0
        total_cost_usd      = 0.0

        def _accum(call_result: dict) -> dict:
            """Add token counts from a call_result to running totals, return result payload."""
            nonlocal total_input_tokens, total_output_tokens, total_cost_usd
            total_input_tokens  += call_result.get("input_tokens",  0)
            total_output_tokens += call_result.get("output_tokens", 0)
            total_cost_usd      += call_result.get("cost_usd",      0.0)
            return call_result["result"]

        # ── Step 3: Parse JD with Claude (1 call, always required) ──────────
        jd_profile = _accum(parse_jd(jd["full_text"], model=user_model))

        # ── Step 4: Load base CVs from storage ───────────────────────────────
        cv_list = []
        cv_raw_bytes = {}          # cv_id → raw docx bytes (reused later)
        for cv_id in request.base_cv_ids:
            cv = get_cv(cv_id, user_id)
            if cv:
                cv_data = download_cv("user_cvs", cv["file_path"])
                doc = Document(io.BytesIO(cv_data))
                cv_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                cv_list.append({"name": cv["filename"], "text": cv_text, "id": cv_id})
                cv_raw_bytes[cv_id] = cv_data

        if not cv_list:
            update_generation(generation_id, "failed", error_message="No valid base CVs found")
            return

        # ── Step 5: Extract JD skills locally (no API) ───────────────────────
        jd_skills = extract_jd_skills_locally(jd["full_text"], jd_profile)

        # ── Step 6: Fetch user's bullet library (no API) ─────────────────────
        library_bullets = get_user_library_bullets(user_id)

        # ── Step 7: Match library bullets against JD skills (no API) ─────────
        match_result = match_jd_to_library(library_bullets, jd_skills, jd["full_text"])
        coverage_pct  = match_result["coverage_pct"]
        covered       = match_result["covered_skills"]
        gap_skills    = match_result["gap_skills"]

        print(f"[gen:{generation_id}] library coverage={coverage_pct:.0%} "
              f"covered={len(covered)} gap={len(gap_skills)} "
              f"library_bullets={len(library_bullets)}")

        # ── Step 8: Choose strategy ───────────────────────────────────────────
        strategy = decide_generation_strategy(coverage_pct)

        best_cv        = cv_list[0]
        api_calls_made = 1          # parse_jd always counts
        api_calls_saved = 0
        gap_bullets: list[str] = []
        assembled: dict = {}

        # ── Strategy A: library_only ──────────────────────────────────────────
        if strategy == "library_only":
            assembled = assemble_cv_from_library(
                match_result=match_result,
                jd_profile=jd_profile,
                base_cv_text=best_cv["text"],
            )
            api_calls_saved = 2

        # ── Strategy B: library_plus_api ─────────────────────────────────────
        elif strategy == "library_plus_api":
            gap_bullets = _accum(fill_gap_bullets(
                base_cv_text=best_cv["text"],
                gap_skills=gap_skills,
                jd_profile=jd_profile,
                n_bullets=min(len(gap_skills) * 2, 8),
                model=user_model,
            ))
            assembled = assemble_cv_from_library(
                match_result=match_result,
                jd_profile=jd_profile,
                base_cv_text=best_cv["text"],
                api_bullets=gap_bullets,
            )
            api_calls_made  += 1
            api_calls_saved  = 1

        # ── Strategy C: full_api ─────────────────────────────────────────────
        else:
            scores_data = _accum(score_cvs(cv_list, jd_profile, model=user_model))
            best_idx = max(range(len(scores_data)), key=lambda i: scores_data[i].get("score", 0))
            best_cv  = cv_list[best_idx]
            api_calls_made += 1

            enh = _accum(enhance_cv(best_cv["text"], jd["full_text"], jd_profile, model=user_model))
            api_calls_made += 1

            assembled = {
                "enhanced_title":    enh.get("enhanced_title", "Cloud Professional"),
                "enhanced_summary":  enh.get("enhanced_summary", ""),
                "bullets":           enh.get("enhanced_bullets", []),
                "injected_keywords": enh.get("injected_keywords", []),
                "source": "full_api",
            }
            api_calls_saved = 0

        # ── Step 8: Build output DOCX ─────────────────────────────────────────
        # If the user provided a template CV, use its formatting
        template_bytes: bytes | None = None
        if request.template_cv_id:
            tmpl_cv = get_cv(request.template_cv_id, user_id)
            if tmpl_cv:
                template_bytes = download_cv("user_cvs", tmpl_cv["file_path"])

        if template_bytes:
            # write_assembled_cv_to_docx returns bytes of the formatted DOCX
            output_bytes = write_assembled_cv_to_docx(template_bytes, assembled)
        else:
            # Fall back: load best base CV and patch title + summary + bullets
            doc = Document(io.BytesIO(cv_raw_bytes.get(best_cv["id"], b"")))

            # Apply title
            if doc.paragraphs and assembled.get("enhanced_title"):
                doc.paragraphs[0].clear()
                doc.paragraphs[0].add_run(assembled["enhanced_title"])

            # Apply summary (find first multi-sentence paragraph after title)
            if assembled.get("enhanced_summary"):
                for para in doc.paragraphs[1:8]:
                    if len(para.text) > 60:
                        para.clear()
                        para.add_run(assembled["enhanced_summary"])
                        break

            # Append gap / library bullets as a new section
            all_bullets = assembled.get("bullets", []) + gap_bullets
            if all_bullets:
                from docx.shared import Pt
                doc.add_heading("Key Achievements", level=2)
                for bullet in all_bullets:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(bullet)

            buf = io.BytesIO()
            doc.save(buf)
            output_bytes = buf.getvalue()

        # ── Step 9: Upload to storage ─────────────────────────────────────────
        filename     = f"enhanced_{int(datetime.utcnow().timestamp())}.docx"
        storage_path = upload_cv(user_id, "generated_cvs", filename, output_bytes)

        # ── Step 10: Persist generated CV record ──────────────────────────────
        generated_cv = create_cv(
            user_id=user_id,
            filename=filename,
            file_path=storage_path,
            file_size_bytes=len(output_bytes),
            cv_type="generated",
            role_title=assembled.get("enhanced_title"),
            seniority=jd_profile.get("seniority"),
            industry=jd.get("industry"),
        )

        # ── Step 11: Index the new generated CV into the bullet library ───────
        try:
            new_bullets = extract_bullets_from_docx(
                docx_bytes=output_bytes,
                cv_meta={
                    "cv_id":       generated_cv["id"],
                    "user_id":     user_id,
                    "role_context":assembled.get("enhanced_title"),
                    "seniority":   jd_profile.get("seniority"),
                }
            )
            if new_bullets:
                save_bullets_for_cv(new_bullets)
                print(f"[gen:{generation_id}] indexed {len(new_bullets)} bullets from generated CV")
        except Exception as idx_err:
            print(f"[gen:{generation_id}] bullet indexing failed (non-fatal): {idx_err}")

        # ── Step 12: Increment usage counters for library bullets used ────────
        used_ids = match_result.get("used_bullet_ids", [])
        if used_ids:
            increment_bullet_usage(used_ids)

        # ── Step 13: Save assembly log ────────────────────────────────────────
        save_assembly_log(
            generation_id=generation_id,
            user_id=user_id,
            strategy=strategy,
            jd_skills_required=len(jd_skills),
            skills_from_library=len(covered),
            skills_from_api=len(gap_skills) if strategy != "library_only" else 0,
            coverage_pct=coverage_pct,
            api_calls_made=api_calls_made,
            api_calls_saved=api_calls_saved,
            tokens_used=total_input_tokens + total_output_tokens,
            bullet_ids_used=used_ids,
            model_used=user_model,
            input_tokens=total_input_tokens,
            output_tokens=total_output_tokens,
            estimated_cost_usd=total_cost_usd,
        )

        # ── Step 14: Update generation record ────────────────────────────────
        processing_time = int((datetime.utcnow() - start_time).total_seconds() * 1000)
        update_generation(
            generation_id,
            "success",
            generated_cv_id=generated_cv["id"],
            generated_cv_file_path=storage_path,
            processing_time_ms=processing_time,
        )

        # ── Step 15: Auto-create submission record ───────────────────────────
        # Pull client/vendor details from the JD record so the Kanban board
        # is pre-populated without the recruiter having to enter them again.
        try:
            submission_data = {
                "generation_id": generation_id,
                "cv_id":         generated_cv["id"],
                "jd_id":         request.jd_id,
                "role_title":    assembled.get("enhanced_title") or jd_profile.get("role_title"),
                "vendor_name":   jd.get("vendor_name"),
                "client_name":   jd.get("client_name"),
                "client_email":  jd.get("client_email"),
            }
            create_submission(user_id=user_id, **submission_data)
            print(f"[gen:{generation_id}] submission record auto-created")
        except Exception as sub_err:
            # Non-fatal — generation already succeeded
            print(f"[gen:{generation_id}] submission auto-create failed (non-fatal): {sub_err}")

        # ── Step 16: Activity log + WebSocket broadcast ───────────────────────
        create_activity_log(
            user_id=user_id,
            action_type="generation_complete",
            description=(
                f"Generated CV for {jd_profile.get('role_title', 'Unknown role')} "
                f"via {strategy} (coverage {coverage_pct:.0%})"
            ),
            metadata={
                "generation_id": generation_id,
                "strategy": strategy,
                "coverage_pct": round(coverage_pct, 3),
                "api_calls_made": api_calls_made,
                "api_calls_saved": api_calls_saved,
                "time_ms": processing_time,
            }
        )

        import asyncio
        asyncio.run(manager.broadcast("generation_completed", {
            "user_id":           user_id,
            "generation_id":     generation_id,
            "status":            "success",
            "strategy":          strategy,
            "coverage_pct":      round(coverage_pct * 100, 1),
            "api_calls_saved":   api_calls_saved,
            "processing_time_ms":processing_time,
        }))

    except Exception as e:
        print(f"[gen:{generation_id}] Generation error: {e}")
        update_generation(generation_id, "failed", error_message=str(e))

@app.get("/api/generation/{generation_id}", response_model=GenerationResponse)
async def get_generation_endpoint(generation_id: str, authorization: str = None):
    """Get generation status"""
    try:
        user_id = get_current_user_id(authorization)
        gen = get_generation(generation_id, user_id)

        if not gen:
            raise HTTPException(status_code=404, detail="Generation not found")

        return GenerationResponse(
            id=gen["id"],
            status=gen["status"],
            jd_id=gen["jd_id"],
            base_cv_ids=gen["base_cv_ids"],
            generated_cv_id=gen.get("generated_cv_id"),
            created_at=gen["created_at"],
            completed_at=gen.get("completed_at"),
            processing_time_ms=gen.get("processing_time_ms"),
            error_message=gen.get("error_message")
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/api/generation/list", response_model=GenerationListResponse)
async def list_generations(authorization: str = None):
    """List user's generations"""
    try:
        user_id = get_current_user_id(authorization)
        generations = get_user_generations(user_id)

        return GenerationListResponse(
            generations=[GenerationResponse(
                id=gen["id"],
                status=gen["status"],
                jd_id=gen["jd_id"],
                base_cv_ids=gen["base_cv_ids"],
                generated_cv_id=gen.get("generated_cv_id"),
                created_at=gen["created_at"],
                completed_at=gen.get("completed_at"),
                processing_time_ms=gen.get("processing_time_ms")
            ) for gen in generations],
            count=len(generations)
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ════════════════════════════════════════════════════════════════
# HISTORY ENDPOINTS
# ════════════════════════════════════════════════════════════════

@app.get("/api/history", response_model=ActivityStreamResponse)
async def get_activity_history(authorization: str = None):
    """Get user's activity history"""
    try:
        user_id = get_current_user_id(authorization)
        activities = get_user_activity(user_id)

        return ActivityStreamResponse(
            logs=[ActivityLogResponse(
                id=activity["id"],
                action_type=activity["action_type"],
                description=activity["description"],
                created_at=activity["created_at"],
                success=activity.get("success", True),
                metadata=activity.get("metadata")
            ) for activity in activities],
            count=len(activities)
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ════════════════════════════════════════════════════════════════
# ADMIN ENDPOINTS
# ════════════════════════════════════════════════════════════════

@app.get("/api/admin/users", response_model=UserListResponse)
async def admin_list_users(authorization: str = None):
    """List all users (admin)"""
    try:
        # TODO: Verify admin role
        users = get_all_users()

        return UserListResponse(
            users=[UserResponse(
                id=user["id"],
                email=user["email"],
                full_name=user.get("full_name"),
                role=user.get("role", "user"),
                created_at=user["created_at"],
                last_login=user.get("last_login"),
                is_active=user.get("is_active", True)
            ) for user in users],
            count=len(users)
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/api/admin/users/{user_id}/activity")
async def admin_user_activity(user_id: str, authorization: str = None):
    """Get specific user's activity (admin)"""
    try:
        # TODO: Verify admin role
        activities = get_user_activity(user_id)

        return ActivityStreamResponse(
            logs=[ActivityLogResponse(
                id=activity["id"],
                action_type=activity["action_type"],
                description=activity["description"],
                created_at=activity["created_at"],
                success=activity.get("success", True),
                metadata=activity.get("metadata")
            ) for activity in activities],
            count=len(activities)
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/api/admin/activity")
async def admin_activity_stream(authorization: str = None):
    """Get all activity logs (admin)"""
    try:
        # TODO: Verify admin role
        activities = get_all_activity()

        return ActivityStreamResponse(
            logs=[ActivityLogResponse(
                id=activity["id"],
                action_type=activity["action_type"],
                description=activity["description"],
                created_at=activity["created_at"],
                success=activity.get("success", True),
                metadata=activity.get("metadata")
            ) for activity in activities],
            count=len(activities)
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ─────────────────────────────────────────────────────────────────────────────
# ADMIN GENERATE — admin generates a CV on behalf of any user
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/api/admin/users/{user_id}/cvs")
async def admin_get_user_cvs(user_id: str, authorization: str = None):
    """Get all CVs belonging to a specific user (for admin generate form)."""
    try:
        admin_id = get_current_user_id(authorization)
        admin    = get_user(admin_id)
        if not admin or admin.get("role") != "admin":
            raise HTTPException(status_code=403, detail="Admin access required")
        cvs = get_user_cvs(user_id)
        return {"cvs": cvs, "count": len(cvs)}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/admin/users/{user_id}/jds")
async def admin_get_user_jds(user_id: str, authorization: str = None):
    """Get all job descriptions belonging to a specific user."""
    try:
        admin_id = get_current_user_id(authorization)
        admin    = get_user(admin_id)
        if not admin or admin.get("role") != "admin":
            raise HTTPException(status_code=403, detail="Admin access required")
        jds = get_user_jds(user_id)
        return {"jds": jds, "count": len(jds)}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/admin/generate")
async def admin_generate_cv(body: dict, authorization: str = None):
    """
    Admin triggers CV generation on behalf of a user.

    Request body:
        {
          "target_user_id": "<uuid>",
          "jd_text":        "<full JD>",
          "base_cv_ids":    ["<uuid>", ...],
          "template_cv_id": "<uuid or null>"
        }

    The generation runs exactly as if the target user clicked Generate,
    using their assigned model and their bullet library.
    A JD record is auto-created under the target user.
    """
    try:
        admin_id = get_current_user_id(authorization)
        admin    = get_user(admin_id)
        if not admin or admin.get("role") != "admin":
            raise HTTPException(status_code=403, detail="Admin access required")

        target_user_id = body.get("target_user_id", "").strip()
        jd_text        = body.get("jd_text", "").strip()
        base_cv_ids    = body.get("base_cv_ids", [])
        template_cv_id = body.get("template_cv_id")

        if not target_user_id:
            raise HTTPException(status_code=400, detail="target_user_id is required")
        if not jd_text:
            raise HTTPException(status_code=400, detail="jd_text is required")
        if not base_cv_ids:
            raise HTTPException(status_code=400, detail="base_cv_ids is required")

        target_user = get_user(target_user_id)
        if not target_user:
            raise HTTPException(status_code=404, detail="Target user not found")

        # Create JD under target user
        jd = create_jd(user_id=target_user_id, full_text=jd_text)
        if not jd:
            raise HTTPException(status_code=500, detail="Failed to create JD")

        # Create generation record under target user
        gen = create_generation(
            user_id=target_user_id,
            jd_id=jd["id"],
            base_cv_ids=base_cv_ids,
            template_cv_id=template_cv_id,
            status="processing",
        )
        if not gen:
            raise HTTPException(status_code=500, detail="Failed to create generation record")

        # Build a surrogate request object for process_generation_sync
        class _AdminGenReq:
            def __init__(self):
                self.jd_id          = jd["id"]
                self.base_cv_ids    = base_cv_ids
                self.template_cv_id = template_cv_id

        # Run synchronously (same as user flow)
        process_generation_sync(gen["id"], target_user_id, _AdminGenReq(), authorization)

        # Log admin action
        create_activity_log(
            user_id=admin_id,
            action_type="admin_generate",
            description=f"Admin {admin['email']} generated CV for user {target_user['email']}",
            metadata={
                "target_user_id": target_user_id,
                "generation_id":  gen["id"],
                "base_cv_ids":    base_cv_ids,
            }
        )

        return {
            "generation_id":  gen["id"],
            "target_user_id": target_user_id,
            "status":         "processing",
            "message":        f"Generation started for {target_user['email']}",
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/admin/usage")
async def admin_usage(authorization: str = None):
    """
    Return per-user cost & token usage stats for every account.
    Powers the Usage & Cost tab in the admin dashboard.
    """
    try:
        get_current_user_id(authorization)   # auth check
        rows = get_all_users_usage_stats()
        # Platform totals
        platform = {
            "total_cost_usd":     round(sum(r["total_cost_usd"]   for r in rows), 4),
            "total_tokens":       sum(r["total_tokens"]            for r in rows),
            "total_generations":  sum(r["total_generations"]       for r in rows),
            "total_calls_saved":  sum(r["api_calls_saved"]         for r in rows),
        }
        return {"users": rows, "platform": platform}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/admin/users/{user_id}/usage")
async def admin_user_usage(user_id: str, authorization: str = None):
    """Detailed cost & usage stats for a single user."""
    try:
        get_current_user_id(authorization)
        stats = get_user_usage_stats(user_id)
        user  = get_user(user_id)
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        return {
            "user_id":   user_id,
            "email":     user["email"],
            "full_name": user.get("full_name"),
            **stats
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/admin/stats", response_model=AdminStatsResponse)
async def admin_stats(authorization: str = None):
    """Get admin dashboard statistics"""
    try:
        # TODO: Verify admin role
        stats = get_admin_stats()

        return AdminStatsResponse(
            total_users=stats.get("total_users", 0),
            active_users_today=0,  # TODO: Calculate
            total_generations=stats.get("total_generations", 0),
            avg_generation_time_ms=0.0,  # TODO: Calculate
            total_cvs=stats.get("total_cvs", 0),
            generations_today=0,  # TODO: Calculate
            errors_today=0  # TODO: Calculate
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ── Model management ──────────────────────────────────────────────────────────

@app.get("/api/admin/models")
async def list_available_models(authorization: str = None):
    """
    Return all Claude models the admin can assign to users.
    The frontend uses this list to populate the model selector dropdown.
    """
    try:
        get_current_user_id(authorization)   # just validates the token
        return {"models": AVAILABLE_MODELS, "default": DEFAULT_MODEL}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.patch("/api/admin/users/{user_id}/model")
async def admin_update_user_model(
    user_id: str,
    body: dict,
    authorization: str = None,
):
    """
    Set the Claude model for a specific user (admin only).

    Request body:
        { "model": "claude-sonnet-4-6", "reason": "Senior role, needs better quality" }
    """
    try:
        admin_id = get_current_user_id(authorization)
        admin = get_user(admin_id)
        if not admin or admin.get("role") != "admin":
            raise HTTPException(status_code=403, detail="Admin access required")

        new_model = body.get("model", "").strip()
        reason    = body.get("reason", "")

        valid_ids = {m["id"] for m in AVAILABLE_MODELS}
        if new_model not in valid_ids:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid model '{new_model}'. Valid options: {sorted(valid_ids)}"
            )

        ok = update_user_model(
            user_id=user_id,
            new_model=new_model,
            admin_email=admin["email"],
            reason=reason or None,
        )
        if not ok:
            raise HTTPException(status_code=500, detail="Failed to update model")

        # Broadcast change to admin WebSocket listeners
        await manager.broadcast("user_model_updated", {
            "user_id":   user_id,
            "new_model": new_model,
            "changed_by":admin["email"],
        })

        return {
            "success":   True,
            "user_id":   user_id,
            "new_model": new_model,
            "message":   f"Model updated to {new_model}",
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/admin/users/{user_id}/model")
async def admin_get_user_model(user_id: str, authorization: str = None):
    """Get the current model assignment + change history for a user."""
    try:
        admin_id = get_current_user_id(authorization)
        admin = get_user(admin_id)
        if not admin or admin.get("role") != "admin":
            raise HTTPException(status_code=403, detail="Admin access required")

        user    = get_user(user_id)
        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        history = get_model_change_history(user_id)

        return {
            "user_id":       user_id,
            "email":         user["email"],
            "current_model": user.get("allowed_model", DEFAULT_MODEL),
            "model_label":   user.get("model_label", "Haiku (Fast · Low Cost)"),
            "updated_at":    user.get("model_updated_at"),
            "updated_by":    user.get("model_updated_by"),
            "history":       history,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ════════════════════════════════════════════════════════════════
# BULK GENERATION
# ════════════════════════════════════════════════════════════════

@app.post("/api/cv/bulk-generate", response_model=BulkGenerationResponse)
async def bulk_generate_cv(
    request: BulkGenerationRequest,
    background_tasks: BackgroundTasks,
    authorization: str = None,
):
    """
    Bulk generate: one candidate (1–3 CVs) against up to 20 JDs in a single
    request.  Each JD becomes an individual generation job running in the
    background.  Returns a bulk_job_id so the client can poll progress.
    """
    try:
        user_id = get_current_user_id(authorization)

        # Create the parent bulk job record
        bulk_job = create_bulk_job(
            user_id=user_id,
            base_cv_ids=request.base_cv_ids,
            template_cv_id=request.template_cv_id,
            total_count=len(request.items),
        )
        if not bulk_job:
            raise HTTPException(status_code=500, detail="Failed to create bulk job")

        bulk_job_id = bulk_job["id"]

        # Schedule each JD as a background task
        for item in request.items:
            background_tasks.add_task(
                _process_bulk_item,
                bulk_job_id=bulk_job_id,
                user_id=user_id,
                base_cv_ids=request.base_cv_ids,
                template_cv_id=request.template_cv_id,
                jd_item=item,
                authorization=authorization,
            )

        create_activity_log(
            user_id=user_id,
            action_type="bulk_generate_started",
            description=f"Bulk generation started: {len(request.items)} JDs",
            metadata={"bulk_job_id": bulk_job_id, "jd_count": len(request.items)},
        )

        return BulkGenerationResponse(
            bulk_job_id=bulk_job_id,
            status="pending",
            total_count=len(request.items),
            message=f"Bulk generation started for {len(request.items)} job descriptions.",
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


def _process_bulk_item(
    bulk_job_id: str,
    user_id: str,
    base_cv_ids: list,
    template_cv_id,
    jd_item,          # BulkJDItem
    authorization: str,
):
    """
    Process one item from a bulk generation batch.
    Creates a JD record, runs generation, and updates bulk job counters.
    """
    try:
        # 1. Create JD record with metadata
        jd = create_jd(
            user_id=user_id,
            full_text=jd_item.jd_text,
            role_title=jd_item.role_title,
            vendor_name=jd_item.vendor_name,
            client_name=jd_item.client_name,
            client_email=jd_item.client_email,
            notes=jd_item.notes,
        )
        if not jd:
            update_bulk_job_progress(bulk_job_id, success=False)
            return

        # 2. Create generation record
        gen = create_generation(
            user_id=user_id,
            jd_id=jd["id"],
            base_cv_ids=base_cv_ids,
            template_cv_id=template_cv_id,
            status="pending",
            bulk_job_id=bulk_job_id,
        )
        if not gen:
            update_bulk_job_progress(bulk_job_id, success=False)
            return

        # 3. Build a surrogate request compatible with process_generation_sync
        class _BulkGenReq:
            def __init__(self):
                self.jd_id          = jd["id"]
                self.base_cv_ids    = base_cv_ids
                self.template_cv_id = template_cv_id

        # 4. Run the same generation logic as single generate
        process_generation_sync(gen["id"], user_id, _BulkGenReq(), authorization)

        # 5. Check outcome and update bulk job counters
        from database import get_generation as _get_gen
        result = _get_gen(gen["id"], user_id)
        succeeded = result and result.get("status") == "success"
        update_bulk_job_progress(bulk_job_id, success=succeeded)

    except Exception as e:
        print(f"[bulk:{bulk_job_id}] item error: {e}")
        update_bulk_job_progress(bulk_job_id, success=False)


@app.get("/api/cv/bulk-generate/{bulk_job_id}")
async def get_bulk_job_status(bulk_job_id: str, authorization: str = None):
    """Get status of a bulk generation job."""
    try:
        user_id  = get_current_user_id(authorization)
        bulk_job = get_user_bulk_jobs(user_id)
        job      = next((j for j in bulk_job if j["id"] == bulk_job_id), None)
        if not job:
            raise HTTPException(status_code=404, detail="Bulk job not found")
        return job
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/cv/bulk-jobs")
async def list_bulk_jobs(authorization: str = None):
    """List all bulk generation jobs for the current user."""
    try:
        user_id = get_current_user_id(authorization)
        jobs    = get_user_bulk_jobs(user_id)
        return {"jobs": jobs, "count": len(jobs)}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ════════════════════════════════════════════════════════════════
# SUBMISSIONS (KANBAN TRACKER)
# ════════════════════════════════════════════════════════════════

@app.post("/api/submissions", response_model=SubmissionResponse)
async def create_submission_endpoint(
    request: SubmissionCreateRequest,
    authorization: str = None,
):
    """Manually create a submission record (auto-created on generation success)."""
    try:
        user_id = get_current_user_id(authorization)
        sub = create_submission(
            user_id=user_id,
            generation_id=request.generation_id,
            cv_id=request.cv_id,
            jd_id=request.jd_id,
            candidate_name=request.candidate_name,
            vendor_name=request.vendor_name,
            client_name=request.client_name,
            client_email=request.client_email,
            role_title=request.role_title,
            notes=request.notes,
        )
        if not sub:
            raise HTTPException(status_code=500, detail="Failed to create submission")

        return SubmissionResponse(**sub)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/submissions", response_model=SubmissionListResponse)
async def list_submissions_endpoint(
    status: str = None,
    authorization: str = None,
):
    """List all submissions for the current user, optionally filtered by status."""
    try:
        user_id     = get_current_user_id(authorization)
        submissions = get_user_submissions(user_id, status_filter=status)
        pipeline    = get_submission_pipeline(user_id)

        return SubmissionListResponse(
            submissions=[SubmissionResponse(**s) for s in submissions],
            count=len(submissions),
            pipeline=pipeline,
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/submissions/{submission_id}", response_model=SubmissionResponse)
async def get_submission_endpoint(submission_id: str, authorization: str = None):
    """Get a single submission."""
    try:
        user_id = get_current_user_id(authorization)
        sub = get_submission(submission_id, user_id)
        if not sub:
            raise HTTPException(status_code=404, detail="Submission not found")
        return SubmissionResponse(**sub)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.patch("/api/submissions/{submission_id}", response_model=SubmissionResponse)
async def update_submission_endpoint(
    submission_id: str,
    request: SubmissionUpdateRequest,
    authorization: str = None,
):
    """
    Update a submission's status or notes.
    Moving to 'submitted' automatically sets submitted_at if not provided.
    """
    try:
        user_id = get_current_user_id(authorization)
        sub = get_submission(submission_id, user_id)
        if not sub:
            raise HTTPException(status_code=404, detail="Submission not found")

        updated = update_submission(
            submission_id=submission_id,
            user_id=user_id,
            status=request.status,
            notes=request.notes,
            submitted_at=request.submitted_at,
            follow_up_at=request.follow_up_at,
            interview_at=request.interview_at,
        )
        if not updated:
            raise HTTPException(status_code=500, detail="Failed to update submission")

        create_activity_log(
            user_id=user_id,
            action_type="submission_update",
            description=f"Submission {submission_id} → {request.status or 'notes updated'}",
            metadata={"submission_id": submission_id, "new_status": request.status},
        )

        return SubmissionResponse(**updated)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/api/submissions/{submission_id}")
async def delete_submission_endpoint(submission_id: str, authorization: str = None):
    """Delete a submission."""
    try:
        user_id = get_current_user_id(authorization)
        sub = get_submission(submission_id, user_id)
        if not sub:
            raise HTTPException(status_code=404, detail="Submission not found")

        from database import delete_submission as _delete_sub
        ok = _delete_sub(submission_id, user_id)
        if not ok:
            raise HTTPException(status_code=500, detail="Failed to delete submission")

        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ════════════════════════════════════════════════════════════════
# WEBSOCKET
# ════════════════════════════════════════════════════════════════

@app.websocket("/ws/admin/{user_id}")
async def websocket_endpoint(websocket: WebSocket, user_id: str):
    """WebSocket endpoint for real-time monitoring"""
    await manager.connect(websocket, user_id)
    try:
        while True:
            data = await websocket.receive_text()
            # Echo back or handle commands
            await websocket.send_text(f"Message received: {data}")
    except WebSocketDisconnect:
        manager.disconnect(user_id)

# ════════════════════════════════════════════════════════════════
# RUN
# ════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main:app",
        host=settings.HOST,
        port=settings.PORT,
        reload=settings.ENV == "development"
    )
